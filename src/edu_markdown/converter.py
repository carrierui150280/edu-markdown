from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

import requests
import yaml
from bs4 import BeautifulSoup
from markdownify import markdownify as html_to_markdown
from docx import Document as DocxDocument
from pypdf import PdfReader
from readability import Document


@dataclass
class ConversionResult:
    title: str
    body_markdown: str
    source: str
    source_type: str


def convert_source(source: str, output_path: Optional[Path] = None) -> Path:
    result = _read_source(source)
    destination = output_path or _default_output_path(source, result.title)

    front_matter = {
        "title": result.title,
        "source": result.source,
        "source_type": result.source_type,
        "fetched_at": _utc_now(),
    }

    content = (
        "---\n"
        + yaml.safe_dump(front_matter, sort_keys=False, allow_unicode=True).strip()
        + "\n---\n\n"
        + result.body_markdown.strip()
        + "\n"
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(content, encoding="utf-8")
    return destination


def convert_directory(
    source_dir: str | Path,
    output_dir: Optional[Path] = None,
    recursive: bool = False,
) -> list[Path]:
    source_root = Path(source_dir)
    if not source_root.exists():
        raise FileNotFoundError(f"Source directory not found: {source_root}")
    if not source_root.is_dir():
        raise NotADirectoryError(f"Source path is not a directory: {source_root}")

    pattern = "**/*" if recursive else "*"
    candidates = sorted(path for path in source_root.glob(pattern) if path.is_file())

    results: list[Path] = []
    for path in candidates:
        if not _is_supported_local_file(path):
            continue

        if output_dir:
            relative = path.relative_to(source_root).with_suffix(".md")
            destination = output_dir / relative
        else:
            destination = path.with_suffix(".md")

        results.append(convert_source(str(path), output_path=destination))

    return results


def _read_source(source: str) -> ConversionResult:
    if _looks_like_url(source):
        return _read_url(source)
    return _read_file(Path(source))


def _read_url(url: str) -> ConversionResult:
    response = requests.get(
        url,
        timeout=20,
        headers={"User-Agent": "edu-markdown/0.1 (+https://example.local)"},
    )
    response.raise_for_status()
    html = response.text
    title, markdown = _html_to_article_markdown(html)
    return ConversionResult(
        title=title,
        body_markdown=markdown,
        source=url,
        source_type="url",
    )


def _read_file(path: Path) -> ConversionResult:
    if not path.exists():
        raise FileNotFoundError(f"Source file not found: {path}")

    suffix = path.suffix.lower()

    if suffix == ".md":
        text = path.read_text(encoding="utf-8")
        title = _infer_title_from_markdown(text) or path.stem
        body = text.strip()
        source_type = "markdown_file"
    elif suffix == ".txt":
        text = path.read_text(encoding="utf-8")
        title = path.stem
        body = text.strip()
        source_type = "text_file"
    elif suffix in {".html", ".htm"}:
        text = path.read_text(encoding="utf-8")
        title, body = _html_to_article_markdown(text)
        source_type = "html_file"
    elif suffix == ".docx":
        title, body = _docx_to_markdown(path)
        source_type = "docx_file"
    elif suffix == ".pdf":
        title, body = _pdf_to_markdown(path)
        source_type = "pdf_file"
    else:
        raise ValueError(
            f"Unsupported file type '{suffix or '(none)'}'. "
            "Use .md, .txt, .html, .htm, .docx, or .pdf for now."
        )

    return ConversionResult(
        title=title,
        body_markdown=body,
        source=str(path.resolve()),
        source_type=source_type,
    )


def _html_to_article_markdown(html: str) -> tuple[str, str]:
    doc = Document(html)
    title = doc.short_title() or _title_from_html(html) or "Untitled"
    article_html = doc.summary(html_partial=True)
    markdown = html_to_markdown(article_html, heading_style="ATX")
    return title.strip(), _normalize_markdown(markdown)


def _title_from_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    if soup.title and soup.title.string:
        return soup.title.string.strip()
    first_heading = soup.find(["h1", "h2"])
    return first_heading.get_text(" ", strip=True) if first_heading else ""


def _infer_title_from_markdown(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
    return ""


def _docx_to_markdown(path: Path) -> tuple[str, str]:
    document = DocxDocument(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]

    if not paragraphs:
        return path.stem, ""

    title = paragraphs[0]
    body_lines = [f"# {title}"]
    body_lines.extend(paragraphs[1:])
    return title, _normalize_markdown("\n\n".join(body_lines))


def _pdf_to_markdown(path: Path) -> tuple[str, str]:
    reader = PdfReader(str(path))
    metadata_title = ""
    if reader.metadata and reader.metadata.title:
        metadata_title = str(reader.metadata.title).strip()

    page_texts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        normalized = _normalize_extracted_text(text)
        if normalized:
            page_texts.append(normalized)

    if not page_texts:
        return metadata_title or path.stem, ""

    paragraphs = []
    for page_text in page_texts:
        paragraphs.extend(
            paragraph.strip() for paragraph in page_text.split("\n\n") if paragraph.strip()
        )

    title = metadata_title or (paragraphs[0].splitlines()[0].strip() if paragraphs else path.stem)
    body_lines = [f"# {title}"]
    remaining = paragraphs
    if remaining and title and remaining[0].startswith(title):
        trimmed_first = remaining[0][len(title) :].strip(" :-\u2013\u2014")
        remaining = ([trimmed_first] if trimmed_first else []) + remaining[1:]
    if remaining:
        body_lines.extend(remaining)
    return title, _normalize_markdown("\n\n".join(body_lines))


def _default_output_path(source: str, title: str) -> Path:
    if _looks_like_url(source):
        safe_name = _slugify(title or urlparse(source).netloc or "page")
        return Path.cwd() / f"{safe_name}.md"

    source_path = Path(source)
    return source_path.with_suffix(".md")


def _looks_like_url(value: str) -> bool:
    parsed = urlparse(value)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def _slugify(value: str) -> str:
    cleaned = "".join(ch.lower() if ch.isalnum() else "-" for ch in value)
    pieces = [piece for piece in cleaned.split("-") if piece]
    return "-".join(pieces) or "output"


def _is_supported_local_file(path: Path) -> bool:
    return path.suffix.lower() in {".md", ".txt", ".html", ".htm", ".docx", ".pdf"}


def _normalize_markdown(value: str) -> str:
    lines = value.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    trimmed = [line.rstrip() for line in lines]
    text = "\n".join(trimmed).strip()
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")
    return text


def _normalize_extracted_text(value: str) -> str:
    lines = [line.strip() for line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    blocks: list[str] = []
    current: list[str] = []

    for line in lines:
        if line:
            current.append(line)
            continue
        if current:
            blocks.append(" ".join(current))
            current = []

    if current:
        blocks.append(" ".join(current))

    return "\n\n".join(blocks).strip()


def _utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace(
        "+00:00", "Z"
    )
