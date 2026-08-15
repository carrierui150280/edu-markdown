from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from edu_markdown.converter import convert_directory, convert_source


def test_convert_txt_file(tmp_path: Path) -> None:
    source = tmp_path / "note.txt"
    source.write_text("hello class", encoding="utf-8")

    output = convert_source(str(source))

    content = output.read_text(encoding="utf-8")
    assert "source_type: text_file" in content
    assert "title: note" in content
    assert "hello class" in content


def test_convert_html_file(tmp_path: Path) -> None:
    source = tmp_path / "lesson.html"
    source.write_text(
        "<html><head><title>Lesson One</title></head>"
        "<body><article><h1>Lesson One</h1><p>Body text.</p></article></body></html>",
        encoding="utf-8",
    )

    output = convert_source(str(source))

    content = output.read_text(encoding="utf-8")
    assert "source_type: html_file" in content
    assert "title: Lesson One" in content
    assert "Body text." in content


def test_convert_docx_file(tmp_path: Path) -> None:
    source = tmp_path / "lesson.docx"
    document = DocxDocument()
    document.add_paragraph("Lesson Two")
    document.add_paragraph("Warm-up activity")
    document.save(source)

    output = convert_source(str(source))

    content = output.read_text(encoding="utf-8")
    assert "source_type: docx_file" in content
    assert "title: Lesson Two" in content
    assert "# Lesson Two" in content
    assert "Warm-up activity" in content


def test_convert_directory_writes_to_output_dir(tmp_path: Path) -> None:
    source_dir = tmp_path / "materials"
    nested_dir = source_dir / "week1"
    nested_dir.mkdir(parents=True)

    (source_dir / "note.txt").write_text("plain text note", encoding="utf-8")
    (nested_dir / "page.html").write_text(
        "<html><head><title>Nested Page</title></head><body><p>Nested body.</p></body></html>",
        encoding="utf-8",
    )

    output_dir = tmp_path / "out"
    results = convert_directory(source_dir, output_dir=output_dir, recursive=True)

    assert len(results) == 2
    assert (output_dir / "note.md").exists()
    assert (output_dir / "week1" / "page.md").exists()


def test_convert_pdf_file(tmp_path: Path) -> None:
    source = tmp_path / "lesson.pdf"
    pdf = canvas.Canvas(str(source), pagesize=letter)
    pdf.setTitle("Lesson Three")
    text = pdf.beginText(72, 720)
    text.textLine("Lesson Three")
    text.textLine("")
    text.textLine("Independent reading practice")
    pdf.drawText(text)
    pdf.save()

    output = convert_source(str(source))

    content = output.read_text(encoding="utf-8")
    assert "source_type: pdf_file" in content
    assert "title: Lesson Three" in content
    assert "# Lesson Three" in content
    assert "Independent reading practice" in content
