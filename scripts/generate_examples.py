from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from edu_markdown.converter import convert_directory


ROOT = Path(__file__).resolve().parents[1]
EXAMPLES_DIR = ROOT / "examples"
SOURCE_DIR = EXAMPLES_DIR / "source"
OUTPUT_DIR = EXAMPLES_DIR / "output"


def main() -> None:
    if EXAMPLES_DIR.exists():
        for child in (SOURCE_DIR, OUTPUT_DIR):
            if child.exists():
                shutil.rmtree(child)

    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    (SOURCE_DIR / "nested").mkdir(parents=True, exist_ok=True)

    _write_text_examples()
    _write_docx_example()
    _write_pdf_example()

    results = convert_directory(SOURCE_DIR, output_dir=OUTPUT_DIR, recursive=True)
    print(f"Generated {len(results)} Markdown example file(s)")


def _write_text_examples() -> None:
    (SOURCE_DIR / "lesson-note.txt").write_text(
        "Warm-up question: What makes a story memorable?\n\n"
        "Students list one character and one conflict.",
        encoding="utf-8",
    )
    (SOURCE_DIR / "nested" / "rubric.md").write_text(
        "# Reading Reflection Rubric\n\n"
        "## Focus\n\n"
        "- clear claim\n"
        "- text evidence\n"
        "- personal response\n",
        encoding="utf-8",
    )
    (SOURCE_DIR / "nested" / "article.html").write_text(
        "<html><head><title>Independent Reading Guide</title></head>"
        "<body><article><h1>Independent Reading Guide</h1>"
        "<p>Choose a book at the right difficulty level.</p>"
        "<p>Write down one quote that shows character growth.</p>"
        "</article></body></html>",
        encoding="utf-8",
    )


def _write_docx_example() -> None:
    doc = DocxDocument()
    doc.add_paragraph("Lesson Plan: Character Motivation")
    doc.add_paragraph("Objective: identify what a character wants and why.")
    doc.add_paragraph("Exit ticket: explain one choice using text evidence.")
    doc.save(SOURCE_DIR / "lesson-plan.docx")


def _write_pdf_example() -> None:
    target = SOURCE_DIR / "reading-handout.pdf"
    pdf = canvas.Canvas(str(target), pagesize=letter)
    pdf.setTitle("Reading Handout")
    text = pdf.beginText(72, 720)
    text.textLine("Reading Handout")
    text.textLine("")
    text.textLine("Underline one sentence that reveals the setting.")
    text.textLine("Circle one word that shows the narrator's tone.")
    pdf.drawText(text)
    pdf.save()


if __name__ == "__main__":
    main()
